"""
step13_requirements_docx_writer

Writes a requirements baseline Word document (.docx) from the list of
extracted requirements stored in step_outputs["3"].

Usage:
    from app.engines.e1.step13_requirements_docx_writer import write_requirements_docx

    docx_path = write_requirements_docx(
        requirements=pipeline.step_outputs["3"],
        opportunity_id=opportunity_id,
        project_name=opportunity.project_name or opportunity_id,
        output_dir=Path(settings.upload_dir) / opportunity_id,
    )
    # docx_path is an absolute Path; store str(docx_path) in step_outputs
"""

from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

_FONT = "Calibri"
_TITLE_SIZE = Pt(20)
_SUBTITLE_SIZE = Pt(10)
_HEADING_SIZE = Pt(14)
_BODY_SIZE = Pt(11)
_META_SIZE = Pt(9)
_HEADING_COLOR = RGBColor(0x1F, 0x38, 0x64)  # dark navy
_META_COLOR = RGBColor(0x75, 0x75, 0x75)  # grey


def _run(para, text: str, size: Pt, bold: bool = False,
         color: RGBColor | None = None, italic: bool = False):
    run = para.add_run(text)
    run.font.name = _FONT
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return run


def write_requirements_docx(
    requirements: list[dict],
    opportunity_id: str,
    project_name: str,
    output_dir: Path,
) -> Path:
    """
    Write a requirements baseline DOCX and return its path.

    Args:
        requirements: list of requirement dicts from step_outputs["3"].
                      Expected keys: id, text, classification, confidence,
                      source_file (all optional - falls back gracefully).
        opportunity_id: used in the output filename.
        project_name: shown in the document title.
        output_dir: directory where the file is saved (created if missing).

    Returns:
        Absolute Path to the generated .docx file.
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    out_path = output_dir / f"requirements_{opportunity_id}.docx"

    doc = Document()

    # Set default Normal style
    normal = doc.styles["Normal"]
    normal.font.name = _FONT
    normal.font.size = _BODY_SIZE

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _run(title_para, "Requirements Baseline", _TITLE_SIZE, bold=True, color=_HEADING_COLOR)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    date_str = datetime.now(timezone.utc).strftime("%d %B %Y")
    _run(
        sub_para,
        f"{project_name}  |  Generated {date_str}  |  {len(requirements)} requirement(s)",
        _SUBTITLE_SIZE,
        color=_META_COLOR,
    )

    doc.add_paragraph()  # spacer

    # Group by classification
    order = ["mandatory", "optional", "conditional"]
    groups: dict[str, list[dict]] = {k: [] for k in order}
    other: list[dict] = []

    for req in requirements:
        cls = str(req.get("classification") or "").lower()
        if cls in groups:
            groups[cls].append(req)
        else:
            other.append(req)

    if other:
        groups["other"] = other
        order.append("other")

    # Write each section
    for cls in order:
        reqs = groups.get(cls, [])
        if not reqs:
            continue

        # Section heading
        heading_para = doc.add_paragraph()
        _run(
            heading_para,
            f"{cls.capitalize()} ({len(reqs)})",
            _HEADING_SIZE,
            bold=True,
            color=_HEADING_COLOR,
        )

        # Requirement items
        for i, req in enumerate(reqs, start=1):
            text = str(req.get("text") or req.get("description") or "").strip()
            req_id = str(req.get("id") or req.get("req_id") or f"{cls[:3].upper()}-{i:03d}")
            confidence = req.get("confidence")
            source = str(req.get("source_file") or req.get("source") or "")

            # Main text
            body_para = doc.add_paragraph(style="List Number")
            _run(body_para, text if text else "(no text)", _BODY_SIZE)

            # Meta line
            meta_parts = [f"ID: {req_id}"]
            if confidence is not None:
                meta_parts.append(f"Confidence: {float(confidence) * 100:.0f}%")
            if source:
                meta_parts.append(f"Source: {source}")

            meta_para = doc.add_paragraph()
            meta_para.paragraph_format.left_indent = Pt(18)
            meta_para.paragraph_format.space_after = Pt(4)
            _run(meta_para, "  |  ".join(meta_parts), _META_SIZE, color=_META_COLOR, italic=True)

        doc.add_paragraph()  # section spacer

    doc.save(out_path)
    return out_path
