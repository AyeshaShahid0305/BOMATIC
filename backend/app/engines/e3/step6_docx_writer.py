import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

_OUTPUT_DIR = Path(__file__).parent / "output"
_HEADING_COLOR = RGBColor(0x00, 0x47, 0xAB)  # Cisco blue
_FONT_NAME = "Arial"
_BODY_SIZE = Pt(11)
_HEADING1_SIZE = Pt(16)
_HEADING2_SIZE = Pt(13)


def _set_font(run, size: Pt, bold: bool = False, color: RGBColor | None = None):
    run.font.name = _FONT_NAME
    run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def _add_page_number(doc: Document):
    """Insert a right-aligned page number footer."""
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para.clear()
    run = para.add_run()
    run.font.name = _FONT_NAME
    run.font.size = Pt(9)
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def _add_cover_page(doc: Document, project_name: str, gbb_tier: str, date_str: str):
    doc.add_paragraph()
    doc.add_paragraph()
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("TECHNICAL PROPOSAL")
    _set_font(run, Pt(28), bold=True, color=_HEADING_COLOR)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_para.add_run(project_name)
    _set_font(run, Pt(18), bold=False)

    doc.add_paragraph()

    tier_para = doc.add_paragraph()
    tier_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tier_para.add_run(f"Solution Tier: {gbb_tier.upper()}")
    _set_font(run, Pt(13), bold=True)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(date_str)
    _set_font(run, _BODY_SIZE)

    doc.add_page_break()


def _add_toc(doc: Document, sections: list[dict]):
    heading = doc.add_paragraph("Table of Contents")
    heading.style = "Heading 1"
    for run in heading.runs:
        _set_font(run, _HEADING1_SIZE, bold=True, color=_HEADING_COLOR)

    for s in sections:
        p = doc.add_paragraph()
        run = p.add_run(f"{s['id'] + 1}.  {s['title']}")
        _set_font(run, _BODY_SIZE)

    doc.add_page_break()


def _add_pricing_table(doc: Document, content: str):
    """Parse the pricing table text and render it as a Word table."""
    lines = content.splitlines()

    # Extract header line, separator, item lines, and footer lines
    item_lines = []
    footer_lines = []
    header_done = False
    in_items = False
    in_footer = False

    for line in lines:
        if line.startswith("-" * 20):
            if not header_done:
                header_done = True
                in_items = True
            else:
                in_items = False
                in_footer = True
            continue
        if in_items and line.strip():
            item_lines.append(line)
        elif in_footer and line.strip():
            footer_lines.append(line)

    if not item_lines:
        doc.add_paragraph(content)
        return

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    headers = ["SKU", "Description", "Qty", "Unit Price", "Line Total"]
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0] if hdr_cells[i].paragraphs[0].runs else hdr_cells[i].paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.name = _FONT_NAME
        run.font.size = _BODY_SIZE

    # Parse item lines (fixed-width format from step5)
    sku_re = re.compile(r"^(\S+)\s{1,}(.+?)\s{2,}(\S+)\s+\S+\s+([\d,]+\.\d{2})\s+\S+\s+([\d,]+\.\d{2})$")
    for line in item_lines:
        m = sku_re.match(line.strip())
        if m:
            row_cells = table.add_row().cells
            values = [m.group(1), m.group(2), m.group(3), m.group(4), m.group(5)]
            for i, v in enumerate(values):
                row_cells[i].text = v
                for run in row_cells[i].paragraphs[0].runs:
                    run.font.name = _FONT_NAME
                    run.font.size = _BODY_SIZE
        else:
            row_cells = table.add_row().cells
            row_cells[0].merge(row_cells[4])
            row_cells[0].text = line.strip()

    doc.add_paragraph()
    for fl in footer_lines:
        p = doc.add_paragraph(fl.strip())
        for run in p.runs:
            run.font.name = _FONT_NAME
            run.font.size = _BODY_SIZE


def _add_section(doc: Document, section: dict):
    heading = doc.add_paragraph(section["title"])
    heading.style = "Heading 1"
    for run in heading.runs:
        _set_font(run, _HEADING1_SIZE, bold=True, color=_HEADING_COLOR)

    content: str = section.get("content", "")

    if section["title"] == "Commercial Proposal":
        _add_pricing_table(doc, content)
        return

    for para_text in content.split("\n\n"):
        para_text = para_text.strip()
        if not para_text:
            continue
        if para_text.startswith("[") and para_text.endswith("]"):
            # Placeholder — render in grey italic
            p = doc.add_paragraph()
            run = p.add_run(para_text)
            _set_font(run, _BODY_SIZE)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            run.font.italic = True
        elif "\n" in para_text:
            # Multi-line block (e.g., compliance table) — preserve lines
            for line in para_text.splitlines():
                p = doc.add_paragraph(line)
                for run in p.runs:
                    run.font.name = _FONT_NAME
                    run.font.size = _BODY_SIZE
        else:
            p = doc.add_paragraph(para_text)
            for run in p.runs:
                _set_font(run, _BODY_SIZE)


def write_proposal(
    assembled_sections: list[dict],
    project_name: str,
    gbb_tier: str = "better",
    output_filename: str | None = None,
) -> Path:
    _OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    if output_filename is None:
        safe_name = re.sub(r"[^\w\- ]", "", project_name).strip().replace(" ", "_")
        output_filename = f"{safe_name}_Technical_Proposal.docx"

    out_path = _OUTPUT_DIR / output_filename

    doc = Document()

    # Set default body font via Normal style
    normal = doc.styles["Normal"]
    normal.font.name = _FONT_NAME
    normal.font.size = _BODY_SIZE

    date_str = datetime.now().strftime("%B %Y")
    _add_cover_page(doc, project_name, gbb_tier, date_str)
    _add_toc(doc, assembled_sections)
    _add_page_number(doc)

    for section in assembled_sections:
        if section["id"] == 0:  # Cover Page already rendered
            continue
        _add_section(doc, section)

    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    from app.engines.e3.step1_template_selector import select_template
    from app.engines.e3.step5_assembler import assemble_proposal
    from app.engines.e3.step8_gbb_pricing import calculate_gbb

    sample_e1 = {
        "project_name": "Riyadh Campus Network",
        "requirements": [
            {"text": "48-port PoE+ switches", "category": "mandatory", "compliance_status": "Compliant"},
            {"text": "10G uplinks", "category": "mandatory", "compliance_status": "Compliant"},
        ],
        "legal_traps": ["Liquidated damages: 0.5%/day uncapped"],
        "missing_documents": [],
    }
    sample_e2 = {
        "matched_items": [
            {"sku": "C9300-48P-E", "product_name": "Cisco Catalyst 9300 48-Port PoE+",
             "qty": 10, "unit_price": 7200.0, "line_total": 61200.0},
        ],
        "unmatched_items": [],
        "subtotal": 61200.0, "discount_amount": 9180.0, "total": 52020.0, "currency": "USD",
    }
    gbb = calculate_gbb(52020.0, "better")
    sections = select_template("rfp")
    narratives = {s.id: "[Section content to be completed manually]" for s in sections}
    assembled = assemble_proposal(sections, narratives, sample_e1, sample_e2, gbb)
    out = write_proposal(assembled, "Riyadh Campus Network", gbb_tier="better")
    print(f"Written: {out}")
