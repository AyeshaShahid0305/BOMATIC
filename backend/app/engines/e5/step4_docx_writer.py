import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from .models import DesignDocument, DesignSection

_OUTPUT_DIR = Path(__file__).parent / "output"

_FONT_NAME = "Arial"
_BODY_SIZE = Pt(11)
_HEADING1_SIZE = Pt(15)
_DIVIDER_SIZE = Pt(18)

_HLD_COLOR = RGBColor(0x00, 0x47, 0xAB)       # Cisco blue
_LLD_COLOR = RGBColor(0x00, 0x82, 0x80)        # Dark cyan / teal
_GREY = RGBColor(0x99, 0x99, 0x99)


def _set_font(run, size: Pt, bold: bool = False, color: RGBColor | None = None,
              italic: bool = False):
    run.font.name = _FONT_NAME
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _add_page_number(doc: Document):
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para.clear()
    run = para.add_run()
    run.font.name = _FONT_NAME
    run.font.size = Pt(9)
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def _add_cover_page(doc: Document, design_doc: DesignDocument):
    doc.add_paragraph()
    doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(design_doc.project_name or "Untitled Project")
    _set_font(run, Pt(28), bold=True, color=_HLD_COLOR)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_para.add_run("High Level Design & Low Level Design")
    _set_font(run, Pt(16), bold=False)

    doc.add_paragraph()

    prep_para = doc.add_paragraph()
    prep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = prep_para.add_run("Prepared by: BOMATIC")
    _set_font(run, Pt(13), bold=True)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(date.today().strftime("%B %d, %Y"))
    _set_font(run, _BODY_SIZE)

    doc.add_page_break()


def _add_toc(doc: Document, hld_sections: list[DesignSection], lld_sections: list[DesignSection]):
    toc_heading = doc.add_paragraph("Table of Contents")
    toc_heading.style = "Heading 1"
    for run in toc_heading.runs:
        _set_font(run, _HEADING1_SIZE, bold=True, color=_HLD_COLOR)

    hld_label = doc.add_paragraph()
    run = hld_label.add_run("HLD")
    _set_font(run, _BODY_SIZE, bold=True, color=_HLD_COLOR)

    for s in hld_sections:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(20)
        run = p.add_run(f"{s.order}.  {s.title}")
        _set_font(run, _BODY_SIZE)

    doc.add_paragraph()

    lld_label = doc.add_paragraph()
    run = lld_label.add_run("LLD")
    _set_font(run, _BODY_SIZE, bold=True, color=_LLD_COLOR)

    for s in lld_sections:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(20)
        run = p.add_run(f"{s.order}.  {s.title}")
        _set_font(run, _BODY_SIZE)

    doc.add_page_break()


def _add_part_divider(doc: Document, label: str, color: RGBColor):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label)
    _set_font(run, _DIVIDER_SIZE, bold=True, color=color)
    doc.add_paragraph()


def _add_design_section(doc: Document, section: DesignSection, heading_color: RGBColor):
    heading = doc.add_paragraph(section.title)
    heading.style = "Heading 1"
    for run in heading.runs:
        _set_font(run, _HEADING1_SIZE, bold=True, color=heading_color)

    for para_text in section.content.split("\n\n"):
        para_text = para_text.strip()
        if not para_text:
            continue
        if para_text.startswith("[") and para_text.endswith("]"):
            p = doc.add_paragraph()
            run = p.add_run(para_text)
            _set_font(run, _BODY_SIZE, italic=True, color=_GREY)
        elif "\n" in para_text:
            for line in para_text.splitlines():
                p = doc.add_paragraph(line)
                for run in p.runs:
                    _set_font(run, _BODY_SIZE)
        else:
            p = doc.add_paragraph(para_text)
            for run in p.runs:
                _set_font(run, _BODY_SIZE)


def write_design_document(design_doc: DesignDocument) -> Path:
    _OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    safe_name = re.sub(r'[\\/*?:"<>|]', "_", design_doc.project_name or "Design").strip()
    out_path = _OUTPUT_DIR / f"{safe_name}_HLD_LLD.docx"

    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = _FONT_NAME
    normal.font.size = _BODY_SIZE

    _add_cover_page(doc, design_doc)
    _add_toc(doc, design_doc.hld_sections, design_doc.lld_sections)
    _add_page_number(doc)

    _add_part_divider(doc, "PART 1 — HIGH LEVEL DESIGN", _HLD_COLOR)
    for i, section in enumerate(design_doc.hld_sections):
        _add_design_section(doc, section, _HLD_COLOR)
        if i == len(design_doc.hld_sections) - 1:
            doc.add_page_break()

    _add_part_divider(doc, "PART 2 — LOW LEVEL DESIGN", _LLD_COLOR)
    for section in design_doc.lld_sections:
        _add_design_section(doc, section, _LLD_COLOR)

    doc.save(out_path)
    return out_path
