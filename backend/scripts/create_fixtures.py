"""
Generate synthetic RFP fixture files for OP-2025-154381 (Aramco Storage).
Mimics the real structure: filenames, section headings, requirement language,
standard references, and BoQ format that the E1 pipeline tests against.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

OUT_DIR = os.path.join(os.path.dirname(__file__), "..", "fixtures", "OP-2025-154381")
os.makedirs(OUT_DIR, exist_ok=True)


# ──────────────────────────────────────────────────────────────
# Helper: styled heading
# ──────────────────────────────────────────────────────────────
def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = RGBColor(0x00, 0x32, 0x6e)
    return p


def add_para(doc, text, bold=False):
    p = doc.add_paragraph(text)
    if bold:
        for run in p.runs:
            run.bold = True
    return p


# ──────────────────────────────────────────────────────────────
# FILE 1: Purchase_Requisition.docx
# ──────────────────────────────────────────────────────────────
def create_purchase_requisition():
    doc = Document()

    # Title block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SAUDI ARAMCO — PURCHASE REQUISITION")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x00, 0x32, 0x6e)

    doc.add_paragraph()
    meta = [
        ("Requisition No.", "OP-2025-154381"),
        ("Project Title",   "Enterprise Storage Infrastructure — DMM7 Data Center Modernization"),
        ("Client",          "Saudi Aramco — IT Infrastructure Division"),
        ("Issue Date",      "01-February-2026"),
        ("Revision",        "Rev. B"),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)

    doc.add_paragraph()
    add_heading(doc, "1. Project Overview", level=1)
    add_para(doc,
        "Saudi Aramco requires the supply, installation, commissioning, and maintenance of an enterprise "
        "storage infrastructure solution for the DMM7 Data Center. The proposed solution must be capable of "
        "supporting the company's critical IT workloads including SAP HANA, Oracle databases, and "
        "virtualization platforms. The vendor shall deliver a complete turnkey solution in accordance with "
        "Saudi Aramco engineering standards and applicable cybersecurity frameworks."
    )

    add_heading(doc, "2. Scope of Work", level=1)
    add_para(doc,
        "The vendor shall provide all necessary hardware, software, licensing, installation, configuration, "
        "testing, commissioning, training, and post-implementation support as described in this requisition "
        "and associated annexures. Refer to Annex A — Detailed Technical Specifications for complete "
        "equipment lists and performance requirements."
    )
    add_para(doc,
        "The vendor shall also comply with all requirements outlined in SACS-002 Third Party Cybersecurity "
        "Standard and SACS-012 Cloud and Outsourcing Cybersecurity Standard as issued by Saudi Aramco "
        "Information Security. Copies of these standards are available upon request — see Section 7."
    )

    add_heading(doc, "3. Mandatory Technical Requirements", level=1)
    add_para(doc, "The following requirements are mandatory. Failure to comply with any mandatory requirement "
                  "shall result in disqualification of the bid.", bold=True)
    doc.add_paragraph()

    mandatory_reqs = [
        ("REQ-001", "The vendor shall provide a minimum raw storage capacity of 5 petabytes (PB) with a "
                    "minimum usable capacity of 3.5 PB after RAID overhead."),
        ("REQ-002", "All storage hardware must be new and unused. Refurbished or end-of-life equipment is "
                    "strictly prohibited and will result in automatic disqualification."),
        ("REQ-003", "The proposed solution shall support NVMe-oF (NVMe over Fabrics) connectivity with "
                    "minimum 32Gb Fibre Channel host interfaces."),
        ("REQ-004", "The vendor shall implement end-to-end data encryption at rest using AES-256 and in "
                    "transit using TLS 1.3 or higher. It is mandatory that encryption keys are managed "
                    "through a dedicated Hardware Security Module (HSM)."),
        ("REQ-005", "The solution must achieve a minimum availability of 99.9999% (six nines) for the "
                    "primary storage array, supported by redundant controllers, power supplies, and "
                    "network interfaces."),
        ("REQ-006", "The vendor shall provide 24×7×365 on-site support with a maximum 4-hour hardware "
                    "replacement response time throughout the 5-year support period."),
        ("REQ-007", "All equipment is required to be on Saudi Aramco's Approved Vendor List (AVL). "
                    "Vendors not on the AVL shall not be considered."),
        ("REQ-008", "The vendor shall comply with SACS-002 cybersecurity requirements including but not "
                    "limited to access control, vulnerability management, and incident response. "
                    "Failure to comply with SACS-002 will result in disqualification."),
        ("REQ-009", "The proposed solution shall support IPv6 dual-stack networking. It is mandatory that "
                    "all management interfaces support both IPv4 and IPv6."),
        ("REQ-010", "The vendor must not subcontract any portion of the installation and commissioning "
                    "work without prior written approval from Saudi Aramco Procurement."),
        ("REQ-011", "The storage solution is required to integrate with the existing VMware vSphere 8.0 "
                    "environment and support VMware VAAI (vStorage APIs for Array Integration) primitives."),
        ("REQ-012", "The vendor shall provide a detailed disaster recovery (DR) plan. The solution must "
                    "support synchronous replication to the DR site with a Recovery Point Objective (RPO) "
                    "of zero and a Recovery Time Objective (RTO) not exceeding 15 minutes."),
    ]

    for req_id, req_text in mandatory_reqs:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"[{req_id}] ").bold = True
        p.add_run(req_text)

    doc.add_paragraph()
    add_heading(doc, "4. Optional / Preferred Requirements", level=1)
    add_para(doc, "The following requirements are desirable but not mandatory. Vendors meeting these "
                  "requirements will receive higher evaluation scores.")
    doc.add_paragraph()

    optional_reqs = [
        ("REQ-013", "The solution should support automated storage tiering (auto-tiering) between NVMe, "
                    "SAS SSD, and NL-SAS tiers without manual intervention."),
        ("REQ-014", "It is preferred that the vendor has prior experience delivering storage solutions "
                    "within the oil and gas sector, specifically for SCADA and historian workloads."),
        ("REQ-015", "The solution may optionally include inline data deduplication and compression to "
                    "improve effective storage utilization."),
        ("REQ-016", "Where possible, the proposed solution should leverage existing Fibre Channel SAN "
                    "infrastructure to minimise cabling and infrastructure changes."),
        ("REQ-017", "It is recommended that the vendor provide a proof-of-concept (PoC) demonstration "
                    "of the proposed solution prior to final award."),
    ]

    for req_id, req_text in optional_reqs:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"[{req_id}] ").bold = True
        p.add_run(req_text)

    doc.add_paragraph()
    add_heading(doc, "5. Conditional Requirements", level=1)
    doc.add_paragraph()

    conditional_reqs = [
        ("REQ-018", "If applicable, the vendor shall comply with IKTVA (In-Kingdom Total Value Add) "
                    "programme requirements with a minimum IKTVA score of 30%."),
        ("REQ-019", "Subject to site access approval by Saudi Aramco Facilities, the vendor shall "
                    "complete physical installation within 90 calendar days of purchase order issuance."),
        ("REQ-020", "Unless otherwise specified in writing by Saudi Aramco IT, all documentation "
                    "shall be delivered in English. Arabic translations may be required at Saudi "
                    "Aramco's sole discretion."),
        ("REQ-021", "Additional storage capacity expansion may be requested at Saudi Aramco's sole "
                    "discretion during the support period, subject to agreed unit pricing."),
        ("REQ-022", "Provided that local content requirements are satisfied, the vendor may propose "
                    "a phased delivery schedule for equipment exceeding 20-week lead times."),
    ]

    for req_id, req_text in conditional_reqs:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"[{req_id}] ").bold = True
        p.add_run(req_text)

    doc.add_paragraph()
    add_heading(doc, "6. Applicable Standards and References", level=1)
    add_para(doc,
        "The vendor shall comply with all applicable standards including but not limited to:\n"
        "  • SACS-002 — Saudi Aramco Third Party Cybersecurity Standard (Rev. 4)\n"
        "  • SACS-012 — Cloud and Outsourcing Cybersecurity Standard\n"
        "  • SACS-333 — Data Classification and Handling Standard\n"
        "  • NCA ECC-2:2024 — National Cybersecurity Authority Essential Cybersecurity Controls\n"
        "  • ISO 27001:2022 — Information Security Management System\n"
        "  • NIST SP 800-53 Rev. 5 — Security and Privacy Controls\n"
        "  • IEEE 802.3 — Ethernet Standards\n"
        "  • SAES-T-916 — Saudi Aramco Engineering Standard for Data Center Infrastructure"
    )

    add_heading(doc, "7. Missing Standards Note", level=1)
    add_para(doc,
        "Note: Copies of SACS-002, SACS-012, and SACS-333 are Saudi Aramco proprietary standards "
        "and are not included in this package. Qualified vendors with an active NDA on file may "
        "request copies from the Procurement Contact. As per Annex B — Vendor Qualification "
        "Requirements, vendors must acknowledge receipt of all applicable standards before bid submission."
    )

    add_heading(doc, "8. Evaluation Criteria", level=1)
    add_para(doc,
        "This tender will be evaluated using a sequential envelope methodology:\n\n"
        "Envelope 1 — Administrative (Pass/Fail):\n"
        "  Complete submission of all mandatory documents. Incomplete submissions may be rejected.\n\n"
        "Envelope 2 — Technical (60% weight):\n"
        "  Minimum passing score: 70 points out of 100\n"
        "  Technical architecture: 30 points\n"
        "  Compliance with SACS/NCA standards: 20 points\n"
        "  Vendor experience and references: 10 points\n\n"
        "Envelope 3 — Commercial (40% weight):\n"
        "  Only opened for vendors passing Envelope 2 technical threshold.\n"
        "  Lowest compliant price receives maximum commercial score.\n\n"
        "IKTVA scoring will be applied separately. Minimum acceptable IKTVA score: 30%."
    )

    path = os.path.join(OUT_DIR, "Purchase_Requisition.docx")
    doc.save(path)
    print(f"  Created: {os.path.basename(path)}")


# ──────────────────────────────────────────────────────────────
# FILE 2: Terms_and_Conditions.docx
# ──────────────────────────────────────────────────────────────
def create_terms_and_conditions():
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TERMS AND CONDITIONS OF TENDER\nOP-2025-154381")
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph()
    add_heading(doc, "1. Definitions", level=1)
    add_para(doc,
        "'Company' means Saudi Aramco. 'Vendor' means the entity submitting a bid in response to this "
        "tender. 'Contract' means the agreement formed upon award of this tender. 'SAR' means Saudi "
        "Arabian Riyal. All monetary values are in SAR unless otherwise stated."
    )

    add_heading(doc, "2. Bid Submission Requirements", level=1)
    add_para(doc,
        "2.1 Closing Date: Bids received after the closing date of 15-March-2026 at 14:00 AST "
        "(Arabia Standard Time) will not be considered under any circumstances. No extensions will "
        "be granted.\n\n"
        "2.2 Submission Format: Bids shall be submitted in sealed envelopes clearly labelled with "
        "the opportunity number OP-2025-154381. The naming convention for electronic copies shall "
        "be: [CompanyName]_OP-2025-154381_Technical.pdf for the technical proposal and "
        "[CompanyName]_OP-2025-154381_Commercial.pdf for the commercial proposal.\n\n"
        "2.3 Required Copies: Vendors shall submit one original and three copies of each envelope. "
        "Electronic copies on USB drives are required in addition to hardcopy submission.\n\n"
        "2.4 Language: All submissions shall be in English. Arabic translations of key sections "
        "may be requested at Saudi Aramco's sole discretion."
    )

    add_heading(doc, "3. Bid Bond / Bank Guarantee", level=1)
    add_para(doc,
        "3.1 A bid bond of 2% of the total bid value is required to be submitted with the "
        "technical envelope. Bids submitted without a valid bid bond shall be disqualified.\n\n"
        "3.2 The bid bond shall be issued by a Saudi Arabian bank or an internationally recognised "
        "bank with a branch in the Kingdom of Saudi Arabia, in favour of Saudi Aramco.\n\n"
        "3.3 Bid bond validity must extend at least 30 days beyond the bid validity period "
        "of 180 calendar days from the closing date.\n\n"
        "3.4 Performance bond of 10% of the contract value will be required within 14 calendar "
        "days of contract award."
    )

    add_heading(doc, "4. Pre-Bid Meeting", level=1)
    add_para(doc,
        "A mandatory pre-bid meeting will be held on 20-February-2026 at 10:00 AST at Saudi Aramco "
        "Headquarters, Dhahran. Attendance is mandatory for all prequalified vendors. Vendors who "
        "do not attend the pre-bid meeting may be disqualified.\n\n"
        "Questions and clarifications must be submitted in writing no later than 25-February-2026. "
        "Responses will be circulated to all prequalified vendors."
    )

    add_heading(doc, "5. Eligibility", level=1)
    add_para(doc,
        "5.1 This tender is restricted to prequalified vendors only. Vendors not on Saudi Aramco's "
        "Approved Vendor List (AVL) for the relevant commodity code are not eligible to bid.\n\n"
        "5.2 Vendors must have a minimum of 5 years experience in delivering enterprise storage "
        "solutions of comparable scale and complexity.\n\n"
        "5.3 Vendors shall provide audited financial statements for the last 3 years demonstrating "
        "financial capability to execute the contract.\n\n"
        "5.4 A conflict of interest declaration must be submitted with the administrative envelope. "
        "Failure to submit this declaration will result in disqualification.\n\n"
        "5.5 An anti-bribery and non-collusion undertaking is required to be signed by an authorised "
        "signatory of the bidding company."
    )

    add_heading(doc, "6. Liquidated Damages", level=1)
    add_para(doc,
        "6.1 In the event of delay in delivery or commissioning beyond the agreed project schedule, "
        "liquidated damages of 0.5% of the total contract value per week of delay shall be applied, "
        "up to a maximum of 10% of the total contract value.\n\n"
        "6.2 Saudi Aramco reserves the right to terminate the contract for convenience at any time "
        "with 30 calendar days written notice, without cause and without liability beyond payment "
        "for work completed to the date of termination.\n\n"
        "6.3 Vendor liability under this contract shall not be limited or capped. The vendor "
        "accepts unlimited liability for all direct damages arising from breach of contract."
    )

    add_heading(doc, "7. Intellectual Property", level=1)
    add_para(doc,
        "7.1 All intellectual property, work product, documentation, custom software, scripts, "
        "and configurations developed specifically for this project shall belong to and be assigned "
        "to Saudi Aramco upon delivery and payment.\n\n"
        "7.2 The vendor hereby assigns all rights, title, and interest in such work product to "
        "Saudi Aramco. The vendor retains no rights to use Saudi Aramco work product for any "
        "other customer or purpose."
    )

    add_heading(doc, "8. Insurance Requirements", level=1)
    add_para(doc,
        "The vendor shall maintain throughout the contract period:\n"
        "  • Professional indemnity insurance: minimum SAR 5,000,000 per occurrence\n"
        "  • Public liability insurance: minimum SAR 10,000,000 per occurrence\n"
        "  • Workers compensation as required by Saudi Arabian labour law\n\n"
        "Certificates of insurance must be provided prior to contract commencement."
    )

    add_heading(doc, "9. Governing Law", level=1)
    add_para(doc,
        "This contract shall be governed by and construed in accordance with the laws of the "
        "Kingdom of Saudi Arabia. Any disputes arising from this contract shall be subject to "
        "the exclusive jurisdiction of the courts of the Kingdom of Saudi Arabia."
    )

    path = os.path.join(OUT_DIR, "Terms_and_Conditions.docx")
    doc.save(path)
    print(f"  Created: {os.path.basename(path)}")


# ──────────────────────────────────────────────────────────────
# FILE 3: BOQ_OP-2025-154381.xlsx
# ──────────────────────────────────────────────────────────────
def create_boq():
    wb = openpyxl.Workbook()

    # ── Sheet 1: BoQ ──
    ws = wb.active
    ws.title = "BoQ"

    header_fill = PatternFill("solid", fgColor="00326E")
    header_font = Font(color="FFFFFF", bold=True, size=11)
    alt_fill    = PatternFill("solid", fgColor="EBF0F8")
    thin        = Side(style="thin", color="AAAAAA")
    border      = Border(left=thin, right=thin, top=thin, bottom=thin)
    center      = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left        = Alignment(horizontal="left",   vertical="center", wrap_text=True)

    # Title
    ws.merge_cells("A1:J1")
    ws["A1"] = "BILL OF QUANTITIES — OP-2025-154381 Enterprise Storage Infrastructure"
    ws["A1"].font = Font(bold=True, size=13, color="00326E")
    ws["A1"].alignment = center

    ws.merge_cells("A2:J2")
    ws["A2"] = "Saudi Aramco — IT Infrastructure Division | Rev. B | 01-Feb-2026"
    ws["A2"].alignment = center

    headers = ["Item #", "Category", "Description", "Vendor / Brand",
               "Part Number", "UOM", "Qty", "Unit List Price (SAR)",
               "Discount %", "Extended Net (SAR)"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=4, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = center
        cell.border = border

    rows = [
        # (item, category, description, vendor, part_no, uom, qty, unit_list, discount)
        ("1.0", "Storage Array",    "All-Flash Storage Array — Primary (5PB raw NVMe)",           "Dell EMC",  "XR7000-5P-NVMe",   "EA",  2,   1_850_000, 15),
        ("1.1", "Storage Array",    "Additional NVMe Expansion Shelf — 24-bay",                    "Dell EMC",  "XR7000-EXP-NVMe",  "EA",  8,     145_000, 12),
        ("1.2", "Storage Array",    "32Gb Fibre Channel Host Bus Adapter (dual-port)",              "Broadcom",  "LPe35002-M2",      "EA", 16,      14_500, 10),
        ("2.0", "Networking",       "32-port 32Gb Fibre Channel Director Switch",                   "Cisco",     "DS-C9148V-32IK9",  "EA",  4,     320_000, 18),
        ("2.1", "Networking",       "32Gb Fibre Channel SFP+ Transceiver (SW)",                    "Cisco",     "DS-SFP-FC32G-SW",  "EA", 64,       3_200,  8),
        ("2.2", "Networking",       "25GbE Ethernet SFP28 DAC Cable — 1m",                         "Cisco",     "SFP-H25G-CU1M",    "EA", 32,         350,  5),
        ("3.0", "Security",         "Hardware Security Module — FIPS 140-2 Level 3",               "Thales",    "LN750-005",        "EA",  2,     185_000, 10),
        ("3.1", "Security",         "FortiGate 1100E Next-Gen Firewall (Storage VLAN)",            "Fortinet",  "FG-1100E",         "EA",  2,     420_000, 20),
        ("3.2", "Security",         "FortiGate 1100E Security Subscription — 5yr",                 "Fortinet",  "FC-10-F11HE-809",  "EA",  2,     245_000, 15),
        ("4.0", "Management",       "Storage Management Software — Enterprise License",            "Dell EMC",  "XR7000-MGMT-ENT",  "EA",  1,     380_000, 12),
        ("4.1", "Management",       "VMware vSphere VAAI Plugin License",                          "VMware",    "VS8-VAAI-K9",      "EA",  1,      42_000, 10),
        ("5.0", "Cabling",          "OM4 Fibre LC-LC Duplex Patch Cable — 5m",                    "Generic",   "OM4-LC-5M",        "EA", 128,         180,  0),
        ("6.0", "Services",         "Installation and Commissioning — Primary Site",               "Vendor",    "SVC-INSTALL-PRI",  "LS",  1,     650_000,  5),
        ("6.1", "Services",         "Installation and Commissioning — DR Site",                    "Vendor",    "SVC-INSTALL-DR",   "LS",  1,     420_000,  5),
        ("6.2", "Services",         "5-Year On-Site Support — 24×7×4hr HW Replacement",           "Dell EMC",  "XR7000-SUP-5Y",    "EA",  2,     580_000, 10),
        ("6.3", "Services",         "Knowledge Transfer and Admin Training (5 days)",              "Vendor",    "SVC-TRAINING-5D",  "EA",  1,      85_000,  0),
        ("7.0", "DR Replication",   "Synchronous Replication Licence — per PB",                   "Dell EMC",  "XR7000-REP-1PB",   "EA",  4,     210_000, 12),
        ("8.0", "Power / Rack",     "42U Data Centre Rack with PDU — Dual Feed",                   "APC",       "AR3100-DPDU",      "EA",  4,      45_000,  8),
        ("8.1", "Power / Rack",     "APC Smart-UPS 10kVA — Double Conversion",                    "APC",       "SRT10KXLI",        "EA",  4,      78_000,  8),
    ]

    for i, (item, cat, desc, vendor, part, uom, qty, unit, disc) in enumerate(rows, 5):
        fill = alt_fill if i % 2 == 0 else PatternFill()
        data = [item, cat, desc, vendor, part, uom, qty, unit, disc, f"=G{i}*H{i}*(1-I{i}/100)"]
        for col, val in enumerate(data, 1):
            cell = ws.cell(row=i, column=col, value=val)
            cell.fill = fill
            cell.alignment = left if col == 3 else center
            cell.border = border
            if col in (8, 10):
                cell.number_format = '#,##0.00'

    total_row = 5 + len(rows)
    ws.merge_cells(f"A{total_row}:I{total_row}")
    ws[f"A{total_row}"] = "TOTAL (BEFORE VAT)"
    ws[f"A{total_row}"].font = Font(bold=True)
    ws[f"A{total_row}"].alignment = Alignment(horizontal="right")
    ws[f"J{total_row}"] = f"=SUM(J5:J{total_row-1})"
    ws[f"J{total_row}"].number_format = '#,##0.00'
    ws[f"J{total_row}"].font = Font(bold=True)

    ws.column_dimensions["A"].width = 7
    ws.column_dimensions["B"].width = 16
    ws.column_dimensions["C"].width = 50
    ws.column_dimensions["D"].width = 14
    ws.column_dimensions["E"].width = 22
    ws.column_dimensions["F"].width = 6
    ws.column_dimensions["G"].width = 6
    ws.column_dimensions["H"].width = 20
    ws.column_dimensions["I"].width = 10
    ws.column_dimensions["J"].width = 22

    # ── Sheet 2: Product_Vendors ──
    pv = wb.create_sheet("Product_Vendors")
    pv["A1"] = "Product Category"
    pv["B1"] = "Preferred Vendor"
    pv["C1"] = "Alternative Vendor"
    pv["D1"] = "Status"
    for cell in pv["1:1"]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = center
        cell.border = border

    vendor_list = [
        ("Storage Arrays — All-Flash NVMe",   "Dell EMC",      "Pure Storage / NetApp", "required"),
        ("Fibre Channel Switches",             "Cisco MDS",     "Brocade / IBM",         "preferred"),
        ("Next-Generation Firewall",           "Fortinet",      "Palo Alto Networks",    "preferred"),
        ("Hardware Security Module",           "Thales",        "Entrust nShield",       "required"),
        ("UPS / Power",                        "APC (Schneider)", "Eaton",               "or_equivalent"),
        ("Server Rack",                        "APC",           "Rittal / Legrand",      "or_equivalent"),
        ("HBA / Host Bus Adapters",            "Broadcom/Emulex", "Marvell/QLogic",      "or_equivalent"),
    ]
    for row, (cat, pref, alt, status) in enumerate(vendor_list, 2):
        pv.cell(row=row, column=1, value=cat).border = border
        pv.cell(row=row, column=2, value=pref).border = border
        pv.cell(row=row, column=3, value=alt).border = border
        pv.cell(row=row, column=4, value=status).border = border

    pv.column_dimensions["A"].width = 36
    pv.column_dimensions["B"].width = 24
    pv.column_dimensions["C"].width = 28
    pv.column_dimensions["D"].width = 16

    path = os.path.join(OUT_DIR, "BOQ_OP-2025-154381.xlsx")
    wb.save(path)
    print(f"  Created: {os.path.basename(path)}")


# ──────────────────────────────────────────────────────────────
# FILE 4: SACS-002_Cybersecurity_Requirements.docx
# ──────────────────────────────────────────────────────────────
def create_sacs_002():
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SACS-002 CYBERSECURITY COMPLIANCE REQUIREMENTS\nOP-2025-154381 — Vendor Obligations")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

    doc.add_paragraph()
    add_para(doc,
        "This document outlines the Saudi Aramco cybersecurity requirements applicable to all "
        "third-party vendors supplying IT infrastructure under OP-2025-154381. These requirements "
        "are derived from SACS-002 (Third Party Cybersecurity Standard, Rev. 4) and are aligned "
        "with the NCA ECC-2:2024 (Essential Cybersecurity Controls, Second Edition, 2024) and "
        "ISO 27001:2022 Annex A controls.", bold=False
    )

    add_heading(doc, "1. Access Control Requirements", level=1)
    access_reqs = [
        ("AC-001", "The vendor shall implement role-based access control (RBAC) for all storage "
                   "management interfaces. It is mandatory that administrative access is restricted "
                   "to authorised Saudi Aramco IT personnel only. This control maps to NCA ECC-2:2024 "
                   "control 2-6 (Identity and Access Management) and ISO 27001:2022 Annex A 5.15."),
        ("AC-002", "Multi-factor authentication (MFA) must be enforced for all remote management "
                   "access to storage systems. Password-only authentication shall not be permitted "
                   "for privileged accounts."),
        ("AC-003", "The vendor shall ensure that all default credentials are changed prior to "
                   "go-live. A credential rotation policy must be implemented with a maximum "
                   "90-day rotation cycle for privileged accounts."),
        ("AC-004", "All access to storage management interfaces shall be logged. Logs must be "
                   "retained for a minimum of 12 months and forwarded to Saudi Aramco's SIEM "
                   "platform in real time."),
    ]
    for req_id, text in access_reqs:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"[{req_id}] ").bold = True
        p.add_run(text)

    add_heading(doc, "2. Encryption Requirements", level=1)
    enc_reqs = [
        ("ENC-001", "All data at rest must be encrypted using AES-256 or stronger. The use of "
                    "deprecated algorithms (DES, 3DES, RC4) is strictly prohibited. This is a "
                    "mandatory requirement per SACS-002 Section 4.3."),
        ("ENC-002", "All data in transit must use TLS 1.3 or higher. TLS 1.0 and TLS 1.1 shall "
                    "be disabled on all management and data interfaces. Failure to comply with "
                    "this encryption standard will result in non-acceptance of the solution."),
        ("ENC-003", "Encryption key management shall be performed through a dedicated Hardware "
                    "Security Module (HSM) that is FIPS 140-2 Level 3 certified. Keys shall "
                    "not be stored on the storage array itself."),
    ]
    for req_id, text in enc_reqs:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"[{req_id}] ").bold = True
        p.add_run(text)

    add_heading(doc, "3. Vulnerability Management", level=1)
    vuln_reqs = [
        ("VM-001", "The vendor shall implement anti-virus and anti-malware protection on all "
                   "management workstations with daily signature updates. This maps to "
                   "NCA ECC-2:2024 control 2-10 (Vulnerability Management) and "
                   "ISO 27001:2022 Annex A 8.7 (Protection against malware)."),
        ("VM-002", "A vulnerability assessment must be performed on all proposed systems prior "
                   "to commissioning. Critical and high vulnerabilities (CVSS score ≥ 7.0) "
                   "must be remediated before go-live."),
        ("VM-003", "Security patches must be applied within 30 days of release for critical "
                   "vulnerabilities and within 90 days for high vulnerabilities throughout "
                   "the support period."),
    ]
    for req_id, text in vuln_reqs:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"[{req_id}] ").bold = True
        p.add_run(text)

    add_heading(doc, "4. Network Security", level=1)
    net_reqs = [
        ("NS-001", "The storage solution shall be deployed in a dedicated storage VLAN that is "
                   "logically isolated from corporate user networks. Network segmentation shall "
                   "be enforced via next-generation firewall policies. This maps to "
                   "NCA ECC-2:2024 control 2-5 (Network Security)."),
        ("NS-002", "Storage management traffic shall be on a dedicated out-of-band management "
                   "network segment. In-band management is permitted only if encrypted with TLS 1.3."),
        ("NS-003", "The vendor shall provide network topology diagrams showing all security zones, "
                   "firewall rules, and data flows. These shall be reviewed and approved by Saudi "
                   "Aramco Information Security before commissioning."),
    ]
    for req_id, text in net_reqs:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"[{req_id}] ").bold = True
        p.add_run(text)

    add_heading(doc, "5. Incident Response", level=1)
    ir_reqs = [
        ("IR-001", "The vendor shall have a documented incident response procedure aligned with "
                   "Saudi Aramco's Cybersecurity Incident Response Plan (CSIRP). The vendor is "
                   "required to notify Saudi Aramco within 2 hours of detecting any security incident."),
        ("IR-002", "If applicable, the vendor shall participate in Saudi Aramco tabletop exercises "
                   "and incident response drills as requested."),
    ]
    for req_id, text in ir_reqs:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"[{req_id}] ").bold = True
        p.add_run(text)

    add_heading(doc, "6. Compliance Framework Mapping", level=1)
    add_para(doc,
        "The requirements in this document map to the following compliance frameworks:\n"
        "  • NCA ECC-2:2024 (National Cybersecurity Authority — Essential Cybersecurity Controls)\n"
        "  • ISO 27001:2022 Annex A (Information Security Controls)\n"
        "  • SAMA CSF (Saudi Arabian Monetary Authority Cybersecurity Framework) — where applicable\n"
        "  • NIST SP 800-53 Rev. 5 — AC, IA, SC, SI control families\n\n"
        "Vendors must submit a completed compliance matrix mapping each requirement above to their "
        "proposed solution. The compliance matrix template is provided in Annex C — Compliance "
        "Matrix Template (not included in this package; request from Procurement)."
    )

    path = os.path.join(OUT_DIR, "SACS-002_Cybersecurity_Requirements.docx")
    doc.save(path)
    print(f"  Created: {os.path.basename(path)}")


# ──────────────────────────────────────────────────────────────
# FILE 5: Technical_Evaluation_Questionnaire.xlsx
# ──────────────────────────────────────────────────────────────
def create_eval_questionnaire():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Evaluation Criteria"

    header_fill = PatternFill("solid", fgColor="00326E")
    header_font = Font(color="FFFFFF", bold=True, size=11)
    thin  = Side(style="thin", color="AAAAAA")
    bdr   = Border(left=thin, right=thin, top=thin, bottom=thin)
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left   = Alignment(horizontal="left",   vertical="center", wrap_text=True)

    ws.merge_cells("A1:G1")
    ws["A1"] = "TECHNICAL EVALUATION QUESTIONNAIRE — OP-2025-154381"
    ws["A1"].font = Font(bold=True, size=13, color="00326E")
    ws["A1"].alignment = center

    ws.merge_cells("A2:G2")
    ws["A2"] = (
        "Scoring: Supported (3 pts) | Will be Customised (2 pts) | Not Supported (0 pts) | "
        "Minimum passing score: 70/100 | Envelope 2 — Technical (60% weight)"
    )
    ws["A2"].alignment = center
    ws["A2"].font = Font(italic=True, size=10)

    headers = ["#", "Category", "Requirement", "Weight", "Response\n(S/WC/NS)", "Score", "Vendor Notes"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=4, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = center
        cell.border = bdr

    criteria = [
        ("1",   "Storage Architecture",    "Solution supports NVMe-oF connectivity (32Gb FC)",              10, "S"),
        ("2",   "Storage Architecture",    "Minimum 5PB raw / 3.5PB usable capacity",                        8, "S"),
        ("3",   "Storage Architecture",    "Six-nines availability (99.9999%) with redundant controllers",    8, "S"),
        ("4",   "Storage Architecture",    "Synchronous replication with RPO=0 and RTO ≤ 15 min",            7, "S"),
        ("5",   "Cybersecurity",           "AES-256 encryption at rest with FIPS 140-2 L3 HSM",             10, "S"),
        ("6",   "Cybersecurity",           "TLS 1.3 for all management and data-in-transit interfaces",       8, "S"),
        ("7",   "Cybersecurity",           "RBAC with MFA for all management access",                         7, "S"),
        ("8",   "Cybersecurity",           "SACS-002 compliance declaration provided",                        5, "S"),
        ("9",   "Integration",             "VMware vSphere 8.0 / VAAI integration",                           6, "S"),
        ("10",  "Integration",             "IPv6 dual-stack management interfaces",                            5, "S"),
        ("11",  "Support",                 "24×7×4hr on-site HW replacement — 5-year term",                  8, "S"),
        ("12",  "Experience",              "≥5 completed references for enterprise storage in O&G sector",    6, "WC"),
        ("13",  "IKTVA",                   "IKTVA score ≥ 30% with detailed IKTVA plan",                      5, "WC"),
        ("14",  "Optional",                "Inline deduplication and compression supported",                   3, "S"),
        ("15",  "Optional",                "Automated storage tiering (NVMe / SAS SSD / NL-SAS)",             4, "S"),
    ]

    score_map = {"S": 3, "WC": 2, "NS": 0}
    for row, (num, cat, req, weight, response) in enumerate(criteria, 5):
        score = score_map.get(response, 0)
        vals = [num, cat, req, weight, response, score, ""]
        fill = PatternFill("solid", fgColor="EBF0F8") if row % 2 == 0 else PatternFill()
        for col, val in enumerate(vals, 1):
            cell = ws.cell(row=row, column=col, value=val)
            cell.fill = fill
            cell.alignment = left if col == 3 else center
            cell.border = bdr

    total_row = 5 + len(criteria)
    ws.merge_cells(f"A{total_row}:D{total_row}")
    ws[f"A{total_row}"] = "TOTAL TECHNICAL SCORE"
    ws[f"A{total_row}"].font = Font(bold=True)
    ws[f"A{total_row}"].alignment = Alignment(horizontal="right")
    ws[f"F{total_row}"] = f"=SUM(F5:F{total_row-1})"
    ws[f"F{total_row}"].font = Font(bold=True)

    ws.column_dimensions["A"].width = 5
    ws.column_dimensions["B"].width = 22
    ws.column_dimensions["C"].width = 55
    ws.column_dimensions["D"].width = 8
    ws.column_dimensions["E"].width = 14
    ws.column_dimensions["F"].width = 8
    ws.column_dimensions["G"].width = 30

    path = os.path.join(OUT_DIR, "Technical_Evaluation_Questionnaire.xlsx")
    wb.save(path)
    print(f"  Created: {os.path.basename(path)}")


# ──────────────────────────────────────────────────────────────
# Run all
# ──────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print(f"\nGenerating fixtures → {OUT_DIR}\n")
    create_purchase_requisition()
    create_terms_and_conditions()
    create_boq()
    create_sacs_002()
    create_eval_questionnaire()
    print(f"\nDone. 5 fixture files created.")
